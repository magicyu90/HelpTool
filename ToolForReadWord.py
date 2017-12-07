import docx
import re  # 引入正则表达式

"""
用正则，匹配 多个数字+一个‘、’为开头，中间夹杂“正确答案”关键字，再以换行符或回车为结尾的。先筛选出题
"""
doc = docx.Document('C:\\Users\\shenyu1.NUCTECH\\Desktop\\test - Copy.docx')
paras = [paragraph.text for paragraph in doc.paragraphs]
topics = []

currentIndex = 0

for para in doc.paragraphs:
    a = re.search(r'\d+、', para.text)  # 题干
    b = re.search(r'^\（A\）|\（B\）|\（C\）|\（D\）', para.text)  # 选项
    if a:
        currentIndex = int(a.group(0)[:-1], 10)
        item = {'number': currentIndex, 'topic': '',
                'answers': [], 'right answers': []}
        item['topic'] += (para.text)
        topics.append(item)
    elif b:
        ans = [answer for answer in re.split(
            r'\（A\）|\（B\）|\（C\）|\（D\）', para.text) if answer != '']
        for item in topics:
            if item['number'] == currentIndex:
                item['answers'].extend(ans)
                break
        for run in para.runs:
            color = int('0x' + run.font.color.rgb.__str__(), 16)
            red = (color & 0xff0000) >> 16
            green = (color & 0xff00) >> 8
            blue = (color & 0xff)
            isred = red > green and red > blue
            if isred and (run.text == "A" or run.text == "B" or run.text == "C" or run.text == "D"):
                for item in topics:
                    if item['number'] == currentIndex:
                        item['right answers'].extend(run.text)
print(topics)
